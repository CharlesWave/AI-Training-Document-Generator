import os
import time
import json
import logging
import traceback
import cv2
import uuid
import datetime
import pprint
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, render_template, redirect, url_for, flash, session
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Import the Gemini API service and prompt
from ai_service import generate_training_document
from prompts_OneStage import raw_response_prompt

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("single_stage_app.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

# Set higher log level for HTTP-related modules
logging.getLogger('httpcore').setLevel(logging.INFO)
logging.getLogger('httpx').setLevel(logging.INFO)
logging.getLogger('urllib3').setLevel(logging.INFO)
logging.getLogger('requests').setLevel(logging.INFO)

# Create custom loggers for key events
video_upload_logger = logging.getLogger('video_upload')
gemini_processing_logger = logging.getLogger('gemini_processing')
document_generation_logger = logging.getLogger('document_generation')

# Configure Flask app
app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SCREENSHOTS_FOLDER'] = 'screenshots'
app.config['ALLOWED_EXTENSIONS'] = {'mp4', 'avi', 'mov', 'wmv', 'mkv'}
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max upload
app.config['PERMANENT_SESSION_LIFETIME'] = 3600  # 1 hour

# Add context processor to provide current date/time to all templates
@app.context_processor
def inject_now():
    return {'now': datetime.datetime.now()}

# Ensure required directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['SCREENSHOTS_FOLDER'], exist_ok=True)

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_screenshots(video_path, timestamps, job_id):
    """Extract screenshots from video at given timestamps"""
    logger.info(f"Extracting screenshots for job {job_id} from {video_path}")
    
    screenshot_paths = []
    screenshot_folder = os.path.join(app.config['SCREENSHOTS_FOLDER'], job_id)
    os.makedirs(screenshot_folder, exist_ok=True)
    
    try:
        # Open the video file
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            logger.error(f"Could not open video file {video_path}")
            return []
        
        # Get video properties
        fps = cap.get(cv2.CAP_PROP_FPS)
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = total_frames / fps if fps > 0 else 0
        
        logger.info(f"Video properties: Duration={duration:.2f}s, FPS={fps:.2f}, Frames={total_frames}")
        
        # Process each timestamp list (for each knowledge point)
        for i, timestamp_entry in enumerate(timestamps):
            point_screenshots = []
            
            # Handle different types of timestamp entries
            if isinstance(timestamp_entry, list):
                # It's already a list as expected
                timestamp_list = timestamp_entry
            elif isinstance(timestamp_entry, int) or isinstance(timestamp_entry, float):
                # Convert single number to a list with one item
                timestamp_list = [str(timestamp_entry)]
            elif timestamp_entry == "":
                # Empty string means no timestamps
                timestamp_list = []
            else:
                # Try to handle it as a string
                try:
                    timestamp_list = [str(timestamp_entry)]
                except:
                    logger.error(f"Cannot process timestamp entry: {timestamp_entry}")
                    timestamp_list = []
            
            # Process each timestamp in the list
            for j, timestamp in enumerate(timestamp_list):
                try:
                    # Convert to string if it's a number
                    if isinstance(timestamp, (int, float)):
                        time_in_seconds = float(timestamp)
                    # Parse timestamp string (assuming format like "1:30")
                    elif ':' in str(timestamp):
                        minutes, seconds = str(timestamp).split(':')
                        time_in_seconds = int(minutes) * 60 + float(seconds)
                    else:
                        # If only seconds are provided as string
                        time_in_seconds = float(timestamp)
                    
                    # Skip if timestamp is beyond video duration
                    if duration > 0 and time_in_seconds > duration:
                        logger.warning(f"Timestamp {timestamp} exceeds video duration of {duration:.2f}s")
                        continue
                    
                    # Set the frame position
                    cap.set(cv2.CAP_PROP_POS_MSEC, time_in_seconds * 1000)
                    
                    # Read the frame
                    success, frame = cap.read()
                    if success:
                        screenshot_filename = f"point_{i+1}_screenshot_{j+1}.png"
                        screenshot_path = os.path.join(screenshot_folder, screenshot_filename)
                        
                        # Save the frame
                        cv2.imwrite(screenshot_path, frame)
                        point_screenshots.append(screenshot_path)
                        logger.info(f"Saved screenshot: {screenshot_path}")
                    else:
                        logger.warning(f"Failed to capture screenshot at timestamp {timestamp}")
                
                except Exception as e:
                    logger.error(f"Error processing timestamp {timestamp}: {str(e)}")
                    logger.error(traceback.format_exc())
            
            screenshot_paths.append(point_screenshots)
        
        # Release the video capture
        cap.release()
        
        logger.info(f"Extracted a total of {sum(len(x) for x in screenshot_paths)} screenshots")
        return screenshot_paths
        
    except Exception as e:
        logger.error(f"Error in extract_screenshots: {str(e)}")
        logger.error(traceback.format_exc())
        return []

def generate_docx(job_id, gemini_response, screenshot_paths):
    """Generate a DOCX document with the extracted knowledge points and screenshots"""
    logger.info(f"Generating DOCX document for job {job_id}")
    
    try:
        # Parse the Gemini response
        data = json.loads(gemini_response)
        
        # Create a new document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Training Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        summary = data.get('Summary', {})
        
        # Add each summary component
        if 'Objective' in summary:
            doc.add_paragraph(f"Objective: {summary['Objective']}")
        elif 'Context' in summary:
            doc.add_paragraph(f"Context: {summary['Context']}")
            
        if 'Key Topics' in summary:
            doc.add_paragraph("Key Topics:")
            topics = summary['Key Topics']
            for topic in (topics if isinstance(topics, list) else [topics]):
                doc.add_paragraph(f"• {topic}", style='List Bullet')
        elif 'Key Topic' in summary:
            doc.add_paragraph("Key Topics:")
            topics = summary['Key Topic']
            for topic in (topics if isinstance(topics, list) else [topics]):
                doc.add_paragraph(f"• {topic}", style='List Bullet')
                
        if 'Key Takeaways' in summary:
            doc.add_paragraph(f"Key Takeaways: {summary['Key Takeaways']}")
        elif 'Key Takeaway' in summary:
            doc.add_paragraph(f"Key Takeaways: {summary['Key Takeaway']}")
            
        # Add a page break
        doc.add_page_break()
        
        # Add knowledge points with screenshots
        doc.add_heading('Knowledge Points', level=1)
        
        # Get the knowledge points and timestamps
        knowledge_points = data.get('Knowledge Points', data.get('Details', []))
        timestamps = data.get('Timestamps', [])
        captions = data.get('Captions', [[]] * len(knowledge_points))  # Default empty captions if not provided
        
        # Add each knowledge point with its screenshots
        for i, (point, point_screenshots) in enumerate(zip(knowledge_points, screenshot_paths)):
            # Add the knowledge point as a heading
            doc.add_heading(f"{i+1}. {point}", level=2)
            
            # Add screenshots for this point
            for j, screenshot_path in enumerate(point_screenshots):
                # Add the screenshot
                try:
                    doc.add_picture(screenshot_path, width=Inches(6))
                    
                    # Add caption if available
                    if i < len(captions) and j < len(captions[i]):
                        caption_text = captions[i][j]
                        caption = doc.add_paragraph(caption_text)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j < len(timestamps[i]):
                        # Use timestamp as caption if no specific caption is available
                        caption = doc.add_paragraph(f"Screenshot at {timestamps[i][j]}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                except Exception as e:
                    logger.error(f"Error adding screenshot {screenshot_path}: {str(e)}")
                    doc.add_paragraph(f"[Error loading screenshot: {str(e)}]")
            
            # Add a page break after each point except the last one
            if i < len(knowledge_points) - 1:
                doc.add_page_break()
        
        # Save the document
        output_path = os.path.join('output', f"{job_id}_training_document.docx")
        os.makedirs('output', exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"Document saved to {output_path}")
        return output_path
    
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        logger.error(traceback.format_exc())
        return None

def debug_log_response(data, job_id):
    """Log the structure of the response data for debugging"""
    try:
        logger.info(f"Response structure for job {job_id}:")
        
        # Log summary
        if 'Summary' in data:
            logger.info("Summary: " + json.dumps(data['Summary'], indent=2))
        else:
            logger.info("Summary section not found in response")
        
        # Log knowledge points count
        knowledge_points = data.get('Knowledge Points', data.get('Details', []))
        logger.info(f"Found {len(knowledge_points)} knowledge points")
        
        # Log timestamps structure
        timestamps = data.get('Timestamps', [])
        logger.info(f"Timestamps structure (length: {len(timestamps)}):")
        for i, ts in enumerate(timestamps[:5]):  # Log first 5 only to avoid too much output
            logger.info(f"  Point {i+1}: {type(ts).__name__} = {ts}")
        
        # Log captions if present
        if 'Captions' in data:
            captions = data['Captions']
            logger.info(f"Captions structure (length: {len(captions)}):")
            for i, cap in enumerate(captions[:3]):  # Log first 3 only
                logger.info(f"  Point {i+1}: {type(cap).__name__} = {cap}")
    
    except Exception as e:
        logger.error(f"Error in debug logging: {str(e)}")

@app.route('/')
def index():
    """Render the home page"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload"""
    try:
        # Check if a file was uploaded
        if 'video' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('index'))
            
        file = request.files['video']
        
        # Check if the file is valid
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('index'))
            
        if not allowed_file(file.filename):
            flash('Invalid file type. Allowed types: mp4, avi, mov, wmv, mkv', 'error')
            return redirect(url_for('index'))
        
        # Generate a unique job ID
        job_id = str(uuid.uuid4())
        
        # Save the uploaded file
        filename = secure_filename(file.filename)
        video_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}_{filename}")
        file.save(video_path)
        
        video_upload_logger.info(f"Video uploaded: {filename}, Job ID: {job_id}")
        logger.info(f"File saved to {video_path}")
        
        # Store the job information in session
        session['job_id'] = job_id
        session['video_path'] = video_path
        session['video_name'] = filename
        
        # Redirect to processing page
        return redirect(url_for('process_video', job_id=job_id))
        
    except Exception as e:
        logger.error(f"Error uploading file: {str(e)}")
        logger.error(traceback.format_exc())
        flash(f"Error uploading file: {str(e)}", 'error')
        return redirect(url_for('index'))

@app.route('/process/<job_id>')
def process_video(job_id):
    """Render the processing page and start the video processing"""
    # Verify job ID matches session
    if session.get('job_id') != job_id:
        flash('Invalid job ID', 'error')
        return redirect(url_for('index'))
        
    video_path = session.get('video_path')
    video_name = session.get('video_name')
    
    if not video_path or not os.path.exists(video_path):
        flash('Video file not found', 'error')
        return redirect(url_for('index'))
    
    # Render the processing page
    return render_template('processing.html', 
                          job_id=job_id, 
                          video_name=video_name,
                          message="Starting video processing...")

@app.route('/start_processing/<job_id>')
def start_processing(job_id):
    """API endpoint to start the processing"""
    try:
        # Verify job ID matches session
        if session.get('job_id') != job_id:
            return {"status": "error", "message": "Invalid job ID"}
            
        video_path = session.get('video_path')
        
        if not video_path or not os.path.exists(video_path):
            return {"status": "error", "message": "Video file not found"}
        
        # Call Gemini API to process the video
        logger.info(f"Starting Gemini processing for job {job_id}")
        gemini_processing_logger.info(f"Starting Gemini processing for job {job_id}, video: {video_path}")
        
        # Generate training document using the Gemini API
        gemini_response = generate_training_document(raw_response_prompt, video_path)
        
        # Check if the response is a processing status
        if isinstance(gemini_response, dict) and gemini_response.get('status') == 'processing':
            # Video is still being processed by Gemini
            session['processing_status'] = gemini_response
            logger.info(f"Video still processing on Gemini: {gemini_response}")
            return {"status": "processing", "message": "Video is still being processed by Gemini API"}
        
        # Store the response in session
        session['gemini_response'] = gemini_response
        gemini_processing_logger.info(f"Gemini processing complete for job {job_id}")
        logger.info(f"Gemini processing complete for job {job_id}")
        
        # Process timestamps and extract screenshots
        try:
            data = json.loads(gemini_response)
            
            # Debug log the structure of the response
            debug_log_response(data, job_id)
            
            timestamps = data.get('Timestamps', [])
            
            logger.info(f"Extracting screenshots for {len(timestamps)} knowledge points")
            screenshot_paths = extract_screenshots(video_path, timestamps, job_id)
            
            # Store the screenshot paths in session
            session['screenshot_paths'] = screenshot_paths
            
            # Generate DOCX document
            document_generation_logger.info(f"Starting document generation for job {job_id}")
            document_path = generate_docx(job_id, gemini_response, screenshot_paths)
            
            if document_path:
                # Store the document path in session
                session['document_path'] = document_path
                document_generation_logger.info(f"Document generation complete for job {job_id}")
                
                return {
                    "status": "complete", 
                    "message": "Processing complete",
                    "redirect": url_for('result', job_id=job_id)
                }
            else:
                return {"status": "error", "message": "Error generating document"}
                
        except Exception as e:
            logger.error(f"Error processing timestamps: {str(e)}")
            logger.error(traceback.format_exc())
            return {"status": "error", "message": f"Error processing timestamps: {str(e)}"}
        
    except Exception as e:
        logger.error(f"Error starting processing: {str(e)}")
        logger.error(traceback.format_exc())
        return {"status": "error", "message": f"Error starting processing: {str(e)}"}

@app.route('/check_status/<job_id>')
def check_status(job_id):
    """API endpoint to check the processing status"""
    # Verify job ID matches session
    if session.get('job_id') != job_id:
        return {"status": "error", "message": "Invalid job ID"}
    
    # Check if processing is complete
    if session.get('document_path'):
        return {
            "status": "complete", 
            "message": "Processing complete",
            "redirect": url_for('result', job_id=job_id)
        }
    
    # Check if there was a processing error
    if session.get('processing_error'):
        return {"status": "error", "message": session.get('processing_error')}
    
    # Check if the video is still being processed by Gemini
    processing_status = session.get('processing_status')
    if processing_status and processing_status.get('status') == 'processing':
        # If we have a file name, check its current state
        file_name = processing_status.get('file_name')
        if file_name:
            from ai_service import get_file_state
            file_state = get_file_state(file_name)
            
            if file_state == "ACTIVE":
                # File is ready, retry processing
                session.pop('processing_status', None)
                return {"status": "retry", "message": "File is ready for processing"}
            elif file_state == "FAILED":
                # File processing failed
                error_msg = "Video processing failed on Gemini API"
                session['processing_error'] = error_msg
                return {"status": "error", "message": error_msg}
            else:
                # Still processing
                return {"status": "processing", "message": "Video is still being processed by Gemini API"}
    
    # Still processing
    return {"status": "processing", "message": "Processing in progress..."}

@app.route('/result/<job_id>')
def result(job_id):
    """Display the processing result"""
    # Verify job ID matches session
    if session.get('job_id') != job_id:
        flash('Invalid job ID', 'error')
        return redirect(url_for('index'))
    
    document_path = session.get('document_path')
    if not document_path or not os.path.exists(document_path):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    # Get the document filename
    document_name = os.path.basename(document_path)
    
    return render_template('result.html', 
                          job_id=job_id,
                          document_name=document_name)

@app.route('/download/<job_id>')
def download_document(job_id):
    """Download the generated document"""
    # Verify job ID matches session
    if session.get('job_id') != job_id:
        flash('Invalid job ID', 'error')
        return redirect(url_for('index'))
    
    document_path = session.get('document_path')
    if not document_path or not os.path.exists(document_path):
        flash('Document not found', 'error')
        return redirect(url_for('index'))
    
    from flask import send_file
    return send_file(document_path, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True) 