import os
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure paths
json_file_path = 'uploads/training_document_1743899404.json'
output_docx_path = 'training_document.docx'
log_file_path = 'screenshot_selection_log.txt'

print(f"Starting document generation process...")

# Create a log file to track screenshot selections
with open(log_file_path, 'w', encoding='utf-8') as log_file:
    log_file.write("SCREENSHOT SELECTION STATUS\n")
    log_file.write("=========================\n\n")

    # Load the JSON data
    print(f"Loading JSON data from {json_file_path}...")
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Create new document
    document = Document()

    # Add title
    title = document.add_heading('Training Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Summary section
    document.add_heading('Summary', level=1)
    document.add_paragraph(data['Summary'])
    document.add_paragraph('')  # Add some space

    # Add Knowledge Points section
    document.add_heading('Knowledge Points', level=1)

    # Process each knowledge point
    total_selected = 0
    total_screenshots = 0

    for i, point_data in enumerate(data['Processing Results']):
        point = point_data['point']
        point_number = i + 1
        log_file.write(f"Point {point_number}: {point}\n")
        
        # Add knowledge point as heading
        document.add_heading(f"{point_number}. {point}", level=2)
        
        # Get screenshots for this point
        screenshots = point_data['screenshots']
        selection_results = point_data.get('selection_results', [])
        
        # Log screenshot information
        total_screenshots += len(selection_results)
        point_selected = 0
        
        # Add screenshot selection info table
        document.add_paragraph("Screenshot Selection Information:")
        
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add table header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Screenshot'
        header_cells[1].text = 'Selected'
        header_cells[2].text = 'Reason/Caption'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Populate table with screenshot info
        for j, result in enumerate(selection_results):
            index = result.get('screenshot_index', -1)
            selected = result.get('selected', False)
            caption = result.get('caption_or_reason', '')
            
            if selected:
                point_selected += 1
                total_selected += 1
            
            if 0 <= index < len(screenshots):
                screenshot_path = screenshots[index]
                filename = os.path.basename(screenshot_path)
                
                # Add to table
                row_cells = table.add_row().cells
                row_cells[0].text = filename
                row_cells[1].text = "Selected" if selected else "Not Selected"
                row_cells[2].text = caption
                
                # Log in file
                log_file.write(f"  - Screenshot {j+1}: {filename} - {'Selected' if selected else 'Not Selected'}\n")
                log_file.write(f"    Caption: {caption}\n")
        
        log_file.write(f"  Summary: {point_selected}/{len(selection_results)} screenshots selected\n\n")
        document.add_paragraph('')
        
        # Add selected screenshots with captions
        selected_count = 0
        for result in selection_results:
            if result.get('selected', False):
                index = result.get('screenshot_index', -1)
                caption = result.get('caption_or_reason', '')
                
                if 0 <= index < len(screenshots):
                    screenshot_path = screenshots[index]
                    
                    # Check if file exists
                    if os.path.exists(screenshot_path):
                        # First selected screenshot? Add heading
                        if selected_count == 0:
                            document.add_heading("Selected Screenshots:", level=3)
                        
                        selected_count += 1
                        
                        try:
                            # Add image to document (6 inches wide max)
                            document.add_picture(screenshot_path, width=Inches(6))
                            
                            # Add caption if available
                            if caption:
                                cap = document.add_paragraph(caption)
                                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            document.add_paragraph('')  # Add space
                        except Exception as e:
                            error_msg = f"Error including image {os.path.basename(screenshot_path)}: {str(e)}"
                            print(error_msg)
                            document.add_paragraph(error_msg)
                    else:
                        error_msg = f"Screenshot file not found: {screenshot_path}"
                        print(error_msg)
                        document.add_paragraph(error_msg)
        
        if selected_count == 0:
            document.add_paragraph("(No screenshots selected for this point)")
        
        # Add page break after each point
        document.add_page_break()

    # Final summary in log file
    log_file.write(f"\nFINAL SUMMARY\n")
    log_file.write(f"=============\n")
    log_file.write(f"Total selected screenshots: {total_selected} out of {total_screenshots}\n")

# Save the document
print(f"Saving DOCX file to {output_docx_path}...")
document.save(output_docx_path)
print(f"DOCX file generated successfully!")
print(f"Screenshot selection log written to {log_file_path}")